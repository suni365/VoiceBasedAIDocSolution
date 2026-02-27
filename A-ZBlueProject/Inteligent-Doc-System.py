import streamlit as st
import docx
import os
import time
import speech_recognition as sr
from lxml import etree
from io import BytesIO
from pydub import AudioSegment

# 1. SET PAGE CONFIG (MUST BE FIRST)
st.set_page_config(layout="wide", page_title="AI-Chatbot")

# Install ffmpeg
os.system("apt-get install -y ffmpeg > /dev/null 2>&1")

# --- SAFE IMPORT BLOCK ---
try:
    from utils import (
    authenticate_user, clean_text, handle_conversation, search_in_doc,
    search_web, save_text_response, search_excel, search_pdf,
    get_base64_image, AudioProcessor,
    search_dat, search_large_xml
)
except SyntaxError as e:
    st.error(f"âŒ Syntax Error in utils.py: {e.msg} at line {e.lineno}")
    st.stop()
except Exception as e:
    st.error(f"âŒ Error importing from utils.py: {e}")
    st.stop()

# --------------------------
# ğŸ”‰ Voice File Processor
# --------------------------
def process_uploaded_voice(voice_file):
    import tempfile
    recognizer = sr.Recognizer()
    try:
        suffix = os.path.splitext(voice_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(voice_file.read())
            tmp_path = tmp_file.name

        if suffix == ".m4a":
            wav_path = tmp_path.replace(".m4a", ".wav")
            AudioSegment.from_file(tmp_path, format="m4a").export(wav_path, format="wav")
        else:
            wav_path = tmp_path

        with sr.AudioFile(wav_path) as source:
            audio = recognizer.record(source)
            return recognizer.recognize_google(audio)
    except Exception as e:
        return f"Error: {e}"

# --------------------------
# ğŸ” Authentication Logic
# --------------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.sidebar.title("ğŸ”‘ Login")
    username_input = st.sidebar.text_input("Username:")
    password_input = st.sidebar.text_input("Password:", type="password")
    
    if st.sidebar.button("Login"):
        if username_input and password_input:
            if authenticate_user(username_input, password_input):
                st.session_state.authenticated = True
                st.session_state["logged_in_user"] = username_input
                st.rerun()
            else:
                st.sidebar.error("âŒ Invalid credentials")
        else:
            st.sidebar.warning("Please enter credentials")
    st.stop()

# --------------------------
# âœ… Main Application (Post-Login)
# --------------------------
st.sidebar.image("A-ZBlueProject/AIChatbot.png")
st.title("ğŸ¤– Intelligent AI-Chatbot for document search ğŸ” ")

# Sidebar User Greeting
try:
    img_base = get_base64_image("A-ZBlueProject/suntita.png.jpg")
    st.markdown(f"""<div style="position:fixed;top:50px;right:10px;background:#333;padding:10px;border-radius:10px;color:white;z-index:999;">
        <img src="data:image/png;base64,{img_base}" width="50"><br><b>{st.session_state['logged_in_user']}</b></div>""", unsafe_allow_html=True)
except: 
    pass

# --------------------------
# ğŸ” Main Document & Voice Search
# --------------------------

# st.divider()
# st.header("ğŸ“˜ WORD DocSearch / ğŸ¤Voice Processing")



# # Added unique keys here to prevent duplicate ID errors
# uploaded_file = st.file_uploader("Upload Word Document (.docx)", type="docx", key="doc_search_uploader")
# user_input = st.text_input("Enter keyword or phrase to search in Doc:")
# voice_file = st.file_uploader("OR Upload Voice (.m4a/.wav)", key="voice_search_uploader")

# response = ""

# # 1. Handle Voice Transcription

# if voice_file:
#     with st.spinner("Transcribing..."):
#         voice_text = process_uploaded_voice(voice_file)
#         if voice_text and not voice_text.startswith("Error"):
#             user_input = voice_text
#             st.success(f"Captured: {user_input}")
#         else:
#             st.error(voice_text)

# # 2. Search Logic (Paragraph based)


# if uploaded_file and user_input:
#     doc = docx.Document(uploaded_file)
#     target = user_input.strip().lower()
#     matches = []
    
#     for para in doc.paragraphs:
#         if target in para.text.lower():
#             if para.text.strip():
#                 matches.append(para.text)
    
#     if matches:
#         st.subheader("Document Matches")
#         for m in matches:
#             st.info(m)
#         response = matches[0] # Use first match for AI context
#     else:
#         st.warning("Phrase not found in document.")

# # 3. AI / Web Fallback
# if user_input and not response:
#     with st.spinner("Consulting AI..."):
#         response = handle_conversation(user_input)
#         if not response or "No relevant info" in response:
#             web_res = search_web(user_input)
#             response = "\n\n".join(web_res) if web_res else "No info found."

# if response:
#     st.markdown(f"<div style='background:#f9f9f9;padding:15px;border-left:5px solid #007bff;color:black;'><b>ğŸ¤– AI Response:</b><br>{response}</div>", unsafe_allow_html=True)
#     try:
#         st.video("A-ZBlueProject/fixed_talking_lady.mp4")
#     except: 
#         pass


# --------------------------
# ğŸ” Word Document Search
# --------------------------

st.header("ğŸ“„ğŸ” Document & Voice Search")

uploaded_file = st.file_uploader("Upload Word Document (.docx)", type="docx", key="doc_search_uploader")
user_input = st.text_input("Enter keyword or phrase to search:")
voice_file = st.file_uploader("ğŸ¤ OR Upload Voice (.m4a/.wav)", key="voice_search_uploader")

response = ""
doc_match_found = False

# ğŸ¤ Voice Processing
if voice_file:
    with st.spinner("Transcribing..."):
        voice_text = process_uploaded_voice(voice_file)
        if voice_text and not voice_text.startswith("Error"):
            user_input = voice_text
            st.success(f"Captured: {user_input}")
        else:
            st.error(voice_text)

# ğŸ“„ Document Search
if uploaded_file and user_input:
    doc = docx.Document(uploaded_file)
    target = user_input.strip().lower()
    matches = []

    for para in doc.paragraphs:
        if target in para.text.lower():
            if para.text.strip():
                matches.append(para.text)

    if matches:
        doc_match_found = True
        st.subheader("ğŸ“„ Document Matches")
        for m in matches:
            st.info(m)
        response = matches[0]
    else:
        st.warning("Phrase not found in document. Checking AI...")

# ğŸ¤– AI + ğŸŒ Web Fallback (ONLY if document failed)
if user_input and not doc_match_found:
    with st.spinner("Consulting AI..."):
        response = handle_conversation(user_input)

        if not response:
            web_res = search_web(user_input)
            response = "\n\n".join(web_res) if web_res else "No info found."

# ğŸ¯ Final Response Display
if response:
    st.markdown(
        f"<div style='background:#f9f9f9;padding:15px;border-left:5px solid #007bff;'>"
        f"<b>ğŸ¤– AI Response:</b><br>{response}</div>",
        unsafe_allow_html=True
    )

    try:
        st.video("A-ZBlueProject/fixed_talking_lady.mp4")
    except:
        pass
# --- Optional Sidebar Search (Excel/PDF) ---

st.divider()
st.header("ğŸ“Š Excel / ğŸ“„ PDF Search")

file_type = st.radio("Select File Type:", ["Excel", "PDF"], key="filetype_radio")

uploaded_data_file = st.file_uploader(
    "Upload Excel (.xlsx) or PDF (.pdf)",
    type=["xlsx", "pdf"],
    key="data_file_uploader"
)

data_keyword = st.text_input(
    "Enter keyword to search in file:",
    key="data_keyword"
)

if uploaded_data_file and data_keyword:

    if file_type == "Excel":
        results = search_excel(uploaded_data_file, data_keyword)

        if results is not None and not results.empty:
            st.subheader("Results")
            st.dataframe(results)
        else:
            st.warning("No matching data found.")

    else:
        results = search_pdf(uploaded_data_file, data_keyword)

        if results:
            st.subheader("Results")
            for r in results:
                st.success(r)
        else:
            st.warning("No matching data found.")

st.divider()
st.header("ğŸ§¾ Structured File Search")

col1, col2 = st.columns(2)

# ğŸ§¾ DAT & XML Search
# --------------------------

# st.divider()
# st.header("ğŸ§¾ Structured File Search")

col1, col2 = st.columns(2)

# -----------------------
# ğŸ“‚ DAT FILE SEARCH
# -----------------------
with col1:
    st.subheader("ğŸ“‚ DAT Search")

    if st.checkbox("Show DAT Tools", key="dat_tools"):
        d_file = st.file_uploader("Upload .dat File", type="dat", key="dat_uploader")
        s_seg = st.text_input("Enter Segment (e.g. NM1*87)", key="dat_segment")

        if st.button("Search DAT", key="dat_search_btn"):
            if d_file and s_seg:
                content = d_file.read().decode("utf-8", errors="ignore")
                dat_results = search_dat(content, s_seg)

                if dat_results:
                    st.success(f"Found {len(dat_results)} matches")
                    for line in dat_results:
                        st.code(line)
                else:
                    st.warning("Segment not found.")
            else:
                st.error("Please upload file and enter segment.")


# -----------------------
# ğŸ” XML CONTEXT SEARCH
# -----------------------
with col2:
    st.subheader("ğŸ” XML Context Search")

    x_file = st.file_uploader("Upload .xml File", type="xml", key="xml_uploader")

    if x_file:
        xtag = st.text_input("Source Tag", key="xml_source_tag")
        xval = st.text_input("Source Value", key="xml_source_value")
        xpath = st.text_input("Target Path (Optional)", key="xml_target_path")

        if st.button("Search XML", key="xml_search_btn"):
            if xtag and xval:
                xml_content = x_file.getvalue()
                x_results = search_large_xml(xml_content, xtag, xval, xpath)

                if x_results:
                    st.success(f"Found {len(x_results)} matches")
                    for r in x_results:
                        st.code(r, language="xml")
                else:
                    st.warning("No matching XML context found.")
            else:
                st.error("Source Tag and Source Value required.")






