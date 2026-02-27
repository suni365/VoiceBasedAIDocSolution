import streamlit as st
import docx
import os
import time
import speech_recognition as sr
import xml.etree.ElementTree as ET
from lxml import etree
from io import BytesIO
from pydub import AudioSegment
from streamlit_webrtc import webrtc_streamer, WebRtcMode

# Install ffmpeg for audio processing (Streamlit Cloud environment)
os.system("apt-get install -y ffmpeg > /dev/null 2>&1")

# --- Utils Import ---
try:
    from utils import (
        authenticate_user, clean_text, handle_conversation, search_in_doc,
        search_web, save_text_response, search_excel, search_pdf,
        get_base64_image, AudioProcessor
    )
except Exception as e:
    st.error(f"Error importing from utils.py: {e}")

# --------------------------
# 🔉 Voice File Processor
# --------------------------
# def process_uploaded_voice(voice_file):
#     import tempfile
#     recognizer = sr.Recognizer()
#     recognizer.energy_threshold = 300
#     recognizer.dynamic_energy_threshold = True
#     tmp_path, wav_path = "", ""
#     try:
#         suffix = os.path.splitext(voice_file.name)[1].lower()
#         with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
#             tmp_file.write(voice_file.read())
#             tmp_path = tmp_file.name

#         if suffix == ".m4a":
#             wav_path = tmp_path.replace(".m4a", ".wav")
#             AudioSegment.from_file(tmp_path, format="m4a").export(wav_path, format="wav")
#         else:
#             wav_path = tmp_path

#         with sr.AudioFile(wav_path) as source:
#             audio = recognizer.record(source)
#             text = recognizer.recognize_google(audio)
#         return text
#     except Exception as e:
#         return f"Error: {e}"
#     finally:
#         if tmp_path and os.path.exists(tmp_path): os.remove(tmp_path)
#         if wav_path and wav_path != tmp_path and os.path.exists(wav_path): os.remove(wav_path)

# --------------------------
# 🔧 XML Helpers
# --------------------------

# def process_uploaded_voice(voice_file):
#     import tempfile
#     recognizer = sr.Recognizer()
#     tmp_path, wav_path = "", ""
#     try:
#         suffix = os.path.splitext(voice_file.name)[1].lower()
        
#         # Save uploaded file to a temporary location
#         with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
#             tmp_file.write(voice_file.read())
#             tmp_path = tmp_file.name

#         if suffix == ".m4a":
#             # This is where ffprobe is needed
#             wav_path = tmp_path.replace(".m4a", ".wav")
#             audio = AudioSegment.from_file(tmp_path, format="m4a")
#             audio.export(wav_path, format="wav")
#         else:
#             wav_path = tmp_path

#         # Recognize the speech
#         with sr.AudioFile(wav_path) as source:
#             audio_data = recognizer.record(source)
#             text = recognizer.recognize_google(audio_data)
#         return text

#     except Exception as e:
#         # If ffprobe is missing, we give a clear instruction
#         if "ffprobe" in str(e) or "ffmpeg" in str(e):
#             return "System Error: ffmpeg is not installed. Please add 'ffmpeg' to a packages.txt file in your repository."
#         return f"Error: {e}"
#     finally:
#         # Cleanup files
#         if tmp_path and os.path.exists(tmp_path): os.remove(tmp_path)
#         if wav_path and wav_path != tmp_path and os.path.exists(wav_path): os.remove(wav_path)

import streamlit as st
import docx
import os
import time
import speech_recognition as sr
from lxml import etree
from io import BytesIO
from pydub import AudioSegment

# Install ffmpeg
os.system("apt-get install -y ffmpeg > /dev/null 2>&1")

# --- SAFE IMPORT BLOCK ---
try:
    from utils import (
        authenticate_user, clean_text, handle_conversation, search_in_doc,
        search_web, save_text_response, search_excel, search_pdf,
        get_base64_image, AudioProcessor
    )
except SyntaxError as e:
    st.error(f"❌ Syntax Error in utils.py: {e.msg} at line {e.lineno}")
    st.stop() # Stop the app until utils.py is fixed
except Exception as e:
    st.error(f"❌ Error importing from utils.py: {e}")
    st.stop()

# --------------------------
# 🔉 Voice File Processor
# --------------------------
def process_uploaded_voice(voice_file):
    import tempfile
    recognizer = sr.Recognizer()
    try:
        suffix = os.path.splitext(voice_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(voice_file.read())
            tmp_path = tmp_file.name

        if suffix == ".m4a":
            wav_path = tmp_path.replace(".m4a", ".wav")
            AudioSegment.from_file(tmp_path, format="m4a").export(wav_path, format="wav")
        else:
            wav_path = tmp_path

        with sr.AudioFile(wav_path) as source:
            audio = recognizer.record(source)
            return recognizer.recognize_google(audio)
    except Exception as e:
        return f"Error: {e}"

# --------------------------
# 🔐 Authentication Logic
# --------------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.sidebar.title("🔑 Login")
    username_input = st.sidebar.text_input("Username:")
    password_input = st.sidebar.text_input("Password:", type="password")
    
    if st.sidebar.button("Login"):
        if username_input and password_input:
            if authenticate_user(username_input, password_input):
                st.session_state.authenticated = True
                st.session_state["logged_in_user"] = username_input
                st.rerun()
            else:
                st.sidebar.error("❌ Invalid credentials")
        else:
            st.sidebar.warning("Please enter credentials")
    st.stop()

# --------------------------
# ✅ Main App (Runs only after login)
# --------------------------
st.title("🤖 Intelligent AI-Chatbot")

# 1. Word Document Search (Paragraph Focus)
st.header("📄 Document Search")
uploaded_file = st.file_uploader("Upload Word Document (.docx)", type="docx")
user_input = st.text_input("Enter keywords (e.g. 3 specific words):")

if uploaded_file and user_input:
    doc = docx.Document(uploaded_file)
    target = user_input.strip().lower()
    found = False
    
    for para in doc.paragraphs:
        if target in para.text.lower():
            if para.text.strip():
                st.info(para.text)
                found = True
    
    if not found:
        st.warning("No matching paragraphs found.")

# 2. Voice Search
voice_file = st.file_uploader("Upload Voice (.m4a/.wav)")
if voice_file:
    with st.spinner("Processing Voice..."):
        transcript = process_uploaded_voice(voice_file)
        st.write(f"**You said:** {transcript}")

def strip_namespace(tag):
    return tag.split('}', 1)[1] if '}' in tag else tag

def search_large_xml(xml_content, source_tag, source_value, target_path):
    results = []
    parser = etree.XMLParser(recover=True)
    tree = etree.parse(BytesIO(xml_content), parser)
    root = tree.getroot()
    
    for elem in root.iter():
        tag_name = strip_namespace(elem.tag)
        if tag_name == source_tag and (elem.text or "").strip() == source_value:
            policy_elem = elem
            while policy_elem is not None and strip_namespace(policy_elem.tag) != "PolicyInfo":
                policy_elem = policy_elem.getparent()
            
            if policy_elem is not None:
                if target_path:
                    for t in policy_elem.iter():
                        if strip_namespace(t.tag) == target_path and t.text:
                            results.append(t.text.strip())
                else:
                    results.append(etree.tostring(policy_elem, pretty_print=True, encoding='unicode'))
    return list(set(results))

# --------------------------
# 🎛️ App Config & Auth
# --------------------------
st.set_page_config(layout="wide", page_title="AI-Chatbot")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

# --------------------------
# ✅ Main Application
# --------------------------
st.sidebar.image("A-ZBlueProject/AIChatbot.png")
st.title("🤖 Intelligent AI-Chatbot")

# Sidebar User Greeting
try:
    img_base = get_base64_image("A-ZBlueProject/suntita.png.jpg")
    st.markdown(f"""<div style="position:fixed;top:50px;right:10px;background:#333;padding:10px;border-radius:10px;color:white;z-index:999;">
        <img src="data:image/png;base64,{img_base}" width="50"><br><b>{st.session_state['logged_in_user']}</b></div>""", unsafe_allow_html=True)
except: pass

# --- Sidebar Excel/PDF Search ---
search_option = st.sidebar.radio("File Search:", ["Excel", "PDF"])
side_key = st.sidebar.text_input("Sidebar Keyword")
if st.sidebar.button("Search Sidebar File"):
    # Insert existing search_excel / search_pdf logic here if needed
    pass

# --------------------------
# 📄 Main Document Search (The Request)
# --------------------------
st.header("🔍 Document & Voice Search")
uploaded_file = st.file_uploader("Upload Word Document (.docx)", type="docx")
user_input = st.text_input("Enter keyword or phrase to search in Doc:")
voice_file = st.file_uploader("OR Upload Voice (.m4a/.wav)")

response = ""

# 1. Handle Voice Transcription
if voice_file:
    with st.spinner("Transcribing..."):
        voice_text = process_uploaded_voice(voice_file)
        if voice_text:
            user_input = voice_text
            st.success(f"Captured: {user_input}")

# 2. Search Logic (Paragraph based)
if uploaded_file and user_input:
    doc = docx.Document(uploaded_file)
    target = user_input.strip().lower()
    matches = []
    
    for para in doc.paragraphs:
        if target in para.text.lower():
            if para.text.strip():
                matches.append(para.text)
    
    if matches:
        st.subheader("Document Matches")
        for m in matches:
            st.info(m)
        response = matches[0] # Use first match for AI context
    else:
        st.warning("Phrase not found in document.")

# 3. AI / Web Fallback
if user_input and not response:
    with st.spinner("Consulting AI..."):
        response = handle_conversation(user_input)
        if not response or "No relevant info" in response:
            web_res = search_web(user_input)
            response = "\n\n".join(web_res) if web_res else "No info found."

if response:
    st.markdown(f"<div style='background:#f9f9f9;padding:15px;border-left:5px solid #007bff;'><b>🤖 AI Response:</b><br>{response}</div>", unsafe_allow_html=True)
    try:
        st.video("A-ZBlueProject/fixed_talking_lady.mp4")
    except: pass

# --------------------------
# 📂 Specialized Search (DAT/XML)
# --------------------------
st.divider()
col1, col2 = st.columns(2)

with col1:
    st.subheader("📂 DAT Search")
    if st.checkbox("Show DAT Tools"):
        d_file = st.file_uploader("Upload .dat", type="dat")
        s_seg = st.text_input("Segment (e.g. NM1*87)")
        if st.button("Search DAT") and d_file:
            # Existing DAT logic
            st.write("Searching...")

with col2:
    st.subheader("🔍 XML Context Search")
    x_file = st.file_uploader("Upload .xml", type="xml")
    if x_file:
        xtag = st.text_input("Source Tag")
        xval = st.text_input("Source Value")
        xpath = st.text_input("Target Path (Optional)")
        if st.button("Search XML"):
            x_results = search_large_xml(x_file.getvalue(), xtag, xval, xpath)
            for r in x_results: st.code(r, language="xml")
